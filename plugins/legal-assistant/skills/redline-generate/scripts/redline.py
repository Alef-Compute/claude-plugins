#!/usr/bin/env python3
"""
Redline Generator with True Track Changes (lxml approach)

Generates DOCX files with actual Word track changes (w:ins, w:del) that can be
accepted/rejected in Microsoft Word's Review tab.

Usage:
    python generate_redlines_trackchanges.py --original-docx <orig.docx> --original-md <orig.md> --modified-md <modified.md> --output <output.docx>

Requirements:
    - python-docx package
    - lxml package (included with python-docx)

Example:
    python generate_redlines_trackchanges.py \
        --original-docx contract.docx \
        --original-md contract_orig.md \
        --modified-md contract_work.md \
        --output contract_redlined.docx
"""

import argparse
import difflib
import re
import sys
from pathlib import Path
from datetime import datetime
from copy import deepcopy

try:
    from docx import Document
    from docx.oxml.ns import qn, nsmap
    from docx.oxml import OxmlElement, register_element_cls
    from lxml import etree
    HAS_DEPS = True
except ImportError as e:
    HAS_DEPS = False
    IMPORT_ERROR = str(e)


# Word XML namespaces
WORD_NAMESPACE = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NSMAP = {'w': WORD_NAMESPACE}

# Global revision ID counter
_revision_id = 0


def get_next_revision_id() -> int:
    """Get the next unique revision ID."""
    global _revision_id
    _revision_id += 1
    return _revision_id


def check_dependencies():
    """Check that all required dependencies are available."""
    errors = []
    if not HAS_DEPS:
        errors.append(f"Missing dependency: {IMPORT_ERROR}\nInstall with: pip install python-docx lxml")
    return errors


def read_file(file_path: str) -> str:
    """Read file contents."""
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()


def normalize_text(text: str) -> str:
    """Normalize text for comparison (collapse whitespace, strip)."""
    return ' '.join(text.split())


def find_text_changes(original_md: str, modified_md: str) -> list[tuple[str, str]]:
    """
    Find specific text changes between original and modified markdown.
    Returns list of (old_text, new_text) tuples representing changes.
    """
    changes = []
    
    orig_lines = original_md.splitlines()
    mod_lines = modified_md.splitlines()
    
    matcher = difflib.SequenceMatcher(None, orig_lines, mod_lines)
    
    for op, i1, i2, j1, j2 in matcher.get_opcodes():
        if op == 'replace':
            orig_text = '\n'.join(orig_lines[i1:i2])
            mod_text = '\n'.join(mod_lines[j1:j2])
            word_changes = find_word_changes(orig_text, mod_text)
            changes.extend(word_changes)
        elif op == 'delete':
            deleted_text = '\n'.join(orig_lines[i1:i2])
            if deleted_text.strip():
                changes.append((deleted_text.strip(), ''))
        elif op == 'insert':
            inserted_text = '\n'.join(mod_lines[j1:j2])
            if inserted_text.strip():
                changes.append(('', inserted_text.strip()))
    
    return changes


def find_word_changes(orig_text: str, mod_text: str) -> list[tuple[str, str]]:
    """Find word-level changes between two text blocks."""
    changes = []
    
    orig_words = re.findall(r'\S+', orig_text)
    mod_words = re.findall(r'\S+', mod_text)
    
    matcher = difflib.SequenceMatcher(None, orig_words, mod_words)
    
    for op, i1, i2, j1, j2 in matcher.get_opcodes():
        if op == 'replace':
            old_phrase = ' '.join(orig_words[i1:i2])
            new_phrase = ' '.join(mod_words[j1:j2])
            if old_phrase != new_phrase:
                changes.append((old_phrase, new_phrase))
        elif op == 'delete':
            old_phrase = ' '.join(orig_words[i1:i2])
            if old_phrase.strip():
                changes.append((old_phrase, ''))
        elif op == 'insert':
            new_phrase = ' '.join(mod_words[j1:j2])
            if new_phrase.strip():
                changes.append(('', new_phrase))
    
    return changes


def create_deletion_element(text: str, author: str, date: str) -> etree.Element:
    """
    Create a w:del element for deleted text.
    
    <w:del w:id="1" w:author="Author" w:date="2024-01-01T00:00:00Z">
      <w:r>
        <w:delText>deleted text</w:delText>
      </w:r>
    </w:del>
    """
    del_elem = etree.Element(qn('w:del'), nsmap=NSMAP)
    del_elem.set(qn('w:id'), str(get_next_revision_id()))
    del_elem.set(qn('w:author'), author)
    del_elem.set(qn('w:date'), date)
    
    # Create run element
    run = etree.SubElement(del_elem, qn('w:r'))
    
    # Create delText element
    del_text = etree.SubElement(run, qn('w:delText'))
    del_text.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    del_text.text = text
    
    return del_elem


def create_insertion_element(text: str, author: str, date: str) -> etree.Element:
    """
    Create a w:ins element for inserted text.
    
    <w:ins w:id="2" w:author="Author" w:date="2024-01-01T00:00:00Z">
      <w:r>
        <w:t>inserted text</w:t>
      </w:r>
    </w:ins>
    """
    ins_elem = etree.Element(qn('w:ins'), nsmap=NSMAP)
    ins_elem.set(qn('w:id'), str(get_next_revision_id()))
    ins_elem.set(qn('w:author'), author)
    ins_elem.set(qn('w:date'), date)
    
    # Create run element
    run = etree.SubElement(ins_elem, qn('w:r'))
    
    # Create text element
    t_elem = etree.SubElement(run, qn('w:t'))
    t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t_elem.text = text
    
    return ins_elem


def find_text_in_runs(paragraph, search_text: str) -> tuple[list, int, int] | None:
    """
    Find text across runs in a paragraph.
    Returns (list of run elements, start offset in first run, end offset in last run) or None.
    """
    full_text = ""
    run_info = []  # [(run_element, start_pos, end_pos, text)]
    
    p_elem = paragraph._p
    
    for run in p_elem.findall(qn('w:r')):
        run_text = ""
        for t_elem in run.findall(qn('w:t')):
            if t_elem.text:
                run_text += t_elem.text
        
        start_pos = len(full_text)
        full_text += run_text
        end_pos = len(full_text)
        run_info.append((run, start_pos, end_pos, run_text))
    
    # Search for the text (with flexible whitespace)
    search_pattern = r'\s*'.join(re.escape(word) for word in search_text.split())
    match = re.search(search_pattern, full_text, re.IGNORECASE)
    
    if not match:
        # Try exact match
        idx = full_text.find(search_text)
        if idx == -1:
            return None
        match_start = idx
        match_end = idx + len(search_text)
    else:
        match_start = match.start()
        match_end = match.end()
    
    # Find which runs contain the match
    affected_runs = []
    for run, start_pos, end_pos, run_text in run_info:
        if start_pos < match_end and end_pos > match_start:
            # This run is affected
            run_match_start = max(0, match_start - start_pos)
            run_match_end = min(len(run_text), match_end - start_pos)
            affected_runs.append((run, run_match_start, run_match_end, run_text))
    
    return affected_runs if affected_runs else None


def apply_track_change_to_paragraph(paragraph, old_text: str, new_text: str, author: str, date: str) -> bool:
    """
    Apply track changes to a paragraph by finding old_text and replacing with tracked deletion/insertion.
    Returns True if change was applied.
    """
    result = find_text_in_runs(paragraph, old_text)
    if not result:
        return False
    
    affected_runs = result
    p_elem = paragraph._p
    
    # For simplicity, we'll handle the case where text is in a single run or spans multiple runs
    if len(affected_runs) == 1:
        # Text is within a single run
        run, start_offset, end_offset, run_text = affected_runs[0]
        
        # Split the run into: before, deleted, inserted, after
        before_text = run_text[:start_offset]
        matched_text = run_text[start_offset:end_offset]
        after_text = run_text[end_offset:]
        
        # Get the run's position in parent
        run_index = list(p_elem).index(run)
        
        # Remove the original run
        p_elem.remove(run)
        
        # Insert new elements at the same position
        insert_pos = run_index
        
        # Add "before" text as regular run
        if before_text:
            before_run = etree.Element(qn('w:r'))
            before_t = etree.SubElement(before_run, qn('w:t'))
            before_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            before_t.text = before_text
            p_elem.insert(insert_pos, before_run)
            insert_pos += 1
        
        # Add deletion element
        if old_text:
            del_elem = create_deletion_element(matched_text, author, date)
            p_elem.insert(insert_pos, del_elem)
            insert_pos += 1
        
        # Add insertion element
        if new_text:
            ins_elem = create_insertion_element(new_text, author, date)
            p_elem.insert(insert_pos, ins_elem)
            insert_pos += 1
        
        # Add "after" text as regular run
        if after_text:
            after_run = etree.Element(qn('w:r'))
            after_t = etree.SubElement(after_run, qn('w:t'))
            after_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            after_t.text = after_text
            p_elem.insert(insert_pos, after_run)
        
        return True
    
    else:
        # Text spans multiple runs - more complex handling
        # For now, we'll consolidate and rebuild
        
        first_run, first_start, first_end, first_text = affected_runs[0]
        last_run, last_start, last_end, last_text = affected_runs[-1]
        
        before_text = first_text[:first_start]
        after_text = last_text[last_end:]
        
        # Collect the matched text
        matched_parts = []
        for run, start, end, text in affected_runs:
            matched_parts.append(text[start:end])
        matched_text = ''.join(matched_parts)
        
        # Find position of first affected run
        first_index = list(p_elem).index(first_run)
        
        # Remove all affected runs
        for run, _, _, _ in affected_runs:
            if run.getparent() is not None:
                p_elem.remove(run)
        
        # Insert new elements
        insert_pos = first_index
        
        if before_text:
            before_run = etree.Element(qn('w:r'))
            before_t = etree.SubElement(before_run, qn('w:t'))
            before_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            before_t.text = before_text
            p_elem.insert(insert_pos, before_run)
            insert_pos += 1
        
        if old_text:
            del_elem = create_deletion_element(matched_text, author, date)
            p_elem.insert(insert_pos, del_elem)
            insert_pos += 1
        
        if new_text:
            ins_elem = create_insertion_element(new_text, author, date)
            p_elem.insert(insert_pos, ins_elem)
            insert_pos += 1
        
        if after_text:
            after_run = etree.Element(qn('w:r'))
            after_t = etree.SubElement(after_run, qn('w:t'))
            after_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            after_t.text = after_text
            p_elem.insert(insert_pos, after_run)
        
        return True


def apply_track_changes_to_document(doc: Document, old_text: str, new_text: str, author: str, date: str) -> bool:
    """
    Search entire document for old_text and apply track changes.
    Returns True if at least one replacement was made.
    """
    # Search in main body paragraphs
    for paragraph in doc.paragraphs:
        if apply_track_change_to_paragraph(paragraph, old_text, new_text, author, date):
            return True
    
    # Search in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if apply_track_change_to_paragraph(paragraph, old_text, new_text, author, date):
                        return True
    
    # Search in headers and footers
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for paragraph in header.paragraphs:
                    if apply_track_change_to_paragraph(paragraph, old_text, new_text, author, date):
                        return True
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                for paragraph in footer.paragraphs:
                    if apply_track_change_to_paragraph(paragraph, old_text, new_text, author, date):
                        return True
    
    return False


def enable_track_changes(doc: Document):
    """Enable track changes in document settings."""
    # Access document settings
    settings = doc.settings
    settings_elem = settings.element
    
    # Add trackRevisions element if not present
    track_revisions = settings_elem.find(qn('w:trackRevisions'))
    if track_revisions is None:
        track_revisions = etree.SubElement(settings_elem, qn('w:trackRevisions'))
    
    # Set to true (val="1" or just presence means true)
    track_revisions.set(qn('w:val'), '1')


def generate_redlines(
    original_docx: str,
    original_md: str,
    modified_md: str,
    output_docx: str,
    author: str = "Legal Review"
) -> bool:
    """
    Generate a redlined DOCX with true track changes.
    """
    orig_docx_path = Path(original_docx)
    orig_md_path = Path(original_md)
    mod_md_path = Path(modified_md)
    out_path = Path(output_docx)
    
    # Validate inputs
    if not orig_docx_path.exists():
        print(f"Error: Original DOCX not found: {original_docx}", file=sys.stderr)
        return False
    
    if not orig_md_path.exists():
        print(f"Error: Original markdown not found: {original_md}", file=sys.stderr)
        return False
    
    if not mod_md_path.exists():
        print(f"Error: Modified markdown not found: {modified_md}", file=sys.stderr)
        return False
    
    try:
        print(f"Loading original document: {orig_docx_path.name}")
        
        # Read markdown files
        orig_md_text = read_file(str(orig_md_path))
        mod_md_text = read_file(str(mod_md_path))
        
        # Find changes
        print("Analyzing changes...")
        changes = find_text_changes(orig_md_text, mod_md_text)
        
        if not changes:
            print("No changes detected between original and modified versions.")
            import shutil
            shutil.copy(orig_docx_path, out_path)
            print(f"Copied original to: {output_docx}")
            return True
        
        print(f"Found {len(changes)} change(s) to apply.")
        
        # Open the original DOCX
        doc = Document(str(orig_docx_path))
        
        # Get current timestamp for revisions
        revision_date = datetime.utcnow().strftime('%Y-%m-%dT%H:%M:%SZ')
        
        # Apply each change with track changes
        changes_applied = 0
        for old_text, new_text in changes:
            if old_text:  # There's something to find and replace
                if apply_track_changes_to_document(doc, old_text, new_text, author, revision_date):
                    changes_applied += 1
                    old_preview = old_text[:40] + '...' if len(old_text) > 40 else old_text
                    new_preview = new_text[:40] + '...' if len(new_text) > 40 else new_text
                    print(f"  Applied: '{old_preview}' -> '{new_preview}'")
                else:
                    print(f"  Warning: Could not find '{old_text[:40]}...' in document")
            elif new_text:
                print(f"  Skipped insertion (no context): '{new_text[:40]}...'")
        
        # Update document properties
        doc.core_properties.last_modified_by = author
        
        # Ensure output directory exists
        out_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Save the modified document
        doc.save(str(out_path))
        
        if out_path.exists():
            print(f"\nSuccess! Redlined document saved to: {output_docx}")
            print(f"  Changes applied: {changes_applied}/{len(changes)}")
            print(f"\nTrack changes features:")
            print(f"  - Deletions: Shown as struck-through (Accept/Reject in Word)")
            print(f"  - Insertions: Shown as underlined (Accept/Reject in Word)")
            print(f"  - Author: {author}")
            print(f"  - Date: {revision_date}")
            return True
        else:
            print("Error: Output file was not created", file=sys.stderr)
            return False
            
    except Exception as e:
        print(f"Error: Redline generation failed: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc()
        return False


def main():
    parser = argparse.ArgumentParser(
        description='Generate DOCX with true Word track changes (Accept/Reject capable)'
    )
    parser.add_argument(
        '--original-docx', '-d',
        help='Path to original DOCX file'
    )
    parser.add_argument(
        '--original-md', '-o',
        help='Path to original markdown file (for comparison)'
    )
    parser.add_argument(
        '--modified-md', '-m',
        help='Path to modified markdown file'
    )
    parser.add_argument(
        '--output', '-out',
        help='Path for output redlined DOCX'
    )
    parser.add_argument(
        '--author', '-a',
        default='Legal Review',
        help='Author name for track changes (default: "Legal Review")'
    )
    parser.add_argument(
        '--check-deps',
        action='store_true',
        help='Check dependencies and exit'
    )
    
    args = parser.parse_args()
    
    errors = check_dependencies()
    
    if args.check_deps:
        if errors:
            print("Missing dependencies:")
            for error in errors:
                print(f"\n  - {error}")
            sys.exit(1)
        else:
            print("All dependencies available!")
            sys.exit(0)
    
    if not args.original_docx or not args.original_md or not args.modified_md or not args.output:
        parser.error("--original-docx, --original-md, --modified-md, and --output are required")
    
    if errors:
        print("Error: Missing required dependencies:", file=sys.stderr)
        for error in errors:
            print(f"\n  - {error}", file=sys.stderr)
        sys.exit(1)
    
    success = generate_redlines(
        original_docx=args.original_docx,
        original_md=args.original_md,
        modified_md=args.modified_md,
        output_docx=args.output,
        author=args.author
    )
    
    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()
